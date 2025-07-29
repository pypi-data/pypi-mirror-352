from docx import Document
from docx.shared import Inches

def create_test_document():
    doc = Document()
    
    # Add title
    doc.add_heading('Test Document', 0)
    
    # Add some paragraphs
    doc.add_paragraph('This is a test document to verify the conversion functionality.')
    
    # Add different heading levels
    doc.add_heading('First Level Heading', level=1)
    doc.add_paragraph('Some content under first level heading.')
    
    doc.add_heading('Second Level Heading', level=2)
    doc.add_paragraph('Content under second level heading.')
    
    # Add a nested list
    doc.add_paragraph('Here is a nested list:')
    for i in range(3):
        p = doc.add_paragraph('Level 1 item ' + str(i+1), style='List Bullet')
        for j in range(2):
            p = doc.add_paragraph('Level 2 item ' + str(j+1), style='List Bullet 2')
            for k in range(2):
                doc.add_paragraph('Level 3 item ' + str(k+1), style='List Bullet 3')
    
    # Add a table
    doc.add_heading('Table Example', level=2)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Header 1'
    header_cells[1].text = 'Header 2'
    header_cells[2].text = 'Header 3'
    
    # Add data rows
    data_cells = table.rows[1].cells
    data_cells[0].text = ''  # Empty cell to test INFO table
    data_cells[1].text = 'Data 2'
    data_cells[2].text = 'Data 3'
    
    data_cells = table.rows[2].cells
    data_cells[0].text = 'Data 4'
    data_cells[1].text = 'Data 5'
    data_cells[2].text = 'Data 6'
    
    # Save the document
    doc.save('test_docs/test.docx')

if __name__ == '__main__':
    create_test_document() 