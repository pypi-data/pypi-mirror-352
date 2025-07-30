"""
DOCX processing utilities for the Agents Hub framework.
"""

from typing import Dict, List, Any, Optional, Union, BinaryIO
import io
import docx
import logging

# Initialize logger
logger = logging.getLogger(__name__)


def extract_text_from_docx(
    file_path: Optional[str] = None,
    file_content: Optional[bytes] = None,
    extract_metadata: bool = True,
) -> Dict[str, Any]:
    """
    Extract text and metadata from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
        file_content: Binary content of the DOCX file
        extract_metadata: Whether to extract metadata
        
    Returns:
        Dictionary containing text and metadata
    """
    if not file_path and not file_content:
        raise ValueError("Either file_path or file_content must be provided")
    
    result = {
        "text": "",
        "metadata": {},
        "paragraphs": [],
    }
    
    try:
        # Open the DOCX file
        if file_path:
            doc = docx.Document(file_path)
        else:
            doc = docx.Document(io.BytesIO(file_content))
        
        # Extract metadata if requested
        if extract_metadata:
            core_properties = doc.core_properties
            result["metadata"] = {
                "title": core_properties.title or "",
                "author": core_properties.author or "",
                "subject": core_properties.subject or "",
                "keywords": core_properties.keywords or "",
                "created": str(core_properties.created) if core_properties.created else "",
                "modified": str(core_properties.modified) if core_properties.modified else "",
                "paragraph_count": len(doc.paragraphs),
            }
        
        # Extract text from each paragraph
        all_text = []
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                all_text.append(paragraph.text)
                result["paragraphs"].append({
                    "index": i,
                    "text": paragraph.text,
                })
        
        # Combine all text
        result["text"] = "\n\n".join(all_text)
        
        return result
    
    except Exception as e:
        logger.exception(f"Error extracting text from DOCX: {e}")
        return {
            "text": "",
            "metadata": {},
            "paragraphs": [],
            "error": str(e),
        }
