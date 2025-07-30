"""
Document processing utilities for the Agents Hub framework.

This module provides utilities for processing documents, including:
- Text extraction from various file formats
- Text chunking for processing large documents
"""

import os
import re
from typing import List, Dict, Any, Optional, Literal


def extract_text_from_pdf(file_path: str) -> Dict[str, Any]:
    """
    Extract text from a PDF file.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Dictionary with extracted text and metadata
    """
    try:
        # Lazy import to avoid dependency if not used
        import PyPDF2
        
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            num_pages = len(reader.pages)
            
            text = ""
            for page_num in range(num_pages):
                page = reader.pages[page_num]
                text += page.extract_text() + "\n\n"
            
            return {
                "text": text,
                "metadata": {
                    "source": os.path.basename(file_path),
                    "type": "pdf",
                    "pages": num_pages,
                }
            }
    
    except ImportError:
        return {
            "error": "PyPDF2 is not installed. Install it with 'pip install PyPDF2'."
        }
    
    except Exception as e:
        return {
            "error": f"Failed to extract text from PDF: {str(e)}"
        }


def extract_text_from_docx(file_path: str) -> Dict[str, Any]:
    """
    Extract text from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Dictionary with extracted text and metadata
    """
    try:
        # Lazy import to avoid dependency if not used
        import docx
        
        doc = docx.Document(file_path)
        
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        return {
            "text": text,
            "metadata": {
                "source": os.path.basename(file_path),
                "type": "docx",
                "paragraphs": len(doc.paragraphs),
            }
        }
    
    except ImportError:
        return {
            "error": "python-docx is not installed. Install it with 'pip install python-docx'."
        }
    
    except Exception as e:
        return {
            "error": f"Failed to extract text from DOCX: {str(e)}"
        }


def extract_text_from_txt(file_path: str) -> Dict[str, Any]:
    """
    Extract text from a TXT file.
    
    Args:
        file_path: Path to the TXT file
        
    Returns:
        Dictionary with extracted text and metadata
    """
    try:
        with open(file_path, "r", encoding="utf-8") as file:
            text = file.read()
        
        return {
            "text": text,
            "metadata": {
                "source": os.path.basename(file_path),
                "type": "txt",
                "size": len(text),
            }
        }
    
    except Exception as e:
        return {
            "error": f"Failed to extract text from TXT: {str(e)}"
        }


def extract_text_from_file(file_path: str) -> Dict[str, Any]:
    """
    Extract text from a file based on its extension.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Dictionary with extracted text and metadata
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == ".pdf":
        return extract_text_from_pdf(file_path)
    
    elif file_extension == ".docx":
        return extract_text_from_docx(file_path)
    
    elif file_extension == ".txt":
        return extract_text_from_txt(file_path)
    
    else:
        return {
            "error": f"Unsupported file format: {file_extension}"
        }


def chunk_text_by_character(
    text: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 200
) -> List[str]:
    """
    Split text into chunks by character count.
    
    Args:
        text: Text to split
        chunk_size: Maximum size of each chunk
        chunk_overlap: Overlap between chunks
        
    Returns:
        List of text chunks
    """
    if not text:
        return []
    
    chunks = []
    start = 0
    
    while start < len(text):
        # Get chunk of size chunk_size
        end = start + chunk_size
        
        # If we're at the end of the text, just add the remaining text
        if end >= len(text):
            chunks.append(text[start:])
            break
        
        # Try to find a good breaking point (newline or space)
        # Look for newline first
        newline_pos = text.rfind("\n", start, end)
        
        if newline_pos > start:
            # Found a newline, use it as the breaking point
            chunks.append(text[start:newline_pos + 1])
            start = newline_pos + 1 - chunk_overlap
        else:
            # No newline, look for space
            space_pos = text.rfind(" ", start, end)
            
            if space_pos > start:
                # Found a space, use it as the breaking point
                chunks.append(text[start:space_pos + 1])
                start = space_pos + 1 - chunk_overlap
            else:
                # No good breaking point, just break at chunk_size
                chunks.append(text[start:end])
                start = end - chunk_overlap
        
        # Ensure start doesn't go negative
        start = max(0, start)
    
    return chunks


def chunk_text_by_sentence(
    text: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 200
) -> List[str]:
    """
    Split text into chunks by sentences, trying to keep chunks under chunk_size.
    
    Args:
        text: Text to split
        chunk_size: Maximum size of each chunk
        chunk_overlap: Overlap between chunks
        
    Returns:
        List of text chunks
    """
    if not text:
        return []
    
    # Simple sentence splitting pattern
    sentence_pattern = r'(?<=[.!?])\s+'
    sentences = re.split(sentence_pattern, text)
    
    chunks = []
    current_chunk = ""
    
    for sentence in sentences:
        # If adding this sentence would exceed chunk_size, start a new chunk
        if len(current_chunk) + len(sentence) > chunk_size and current_chunk:
            chunks.append(current_chunk)
            
            # Start new chunk with overlap
            # Find the last few sentences to use as overlap
            overlap_text = ""
            overlap_size = 0
            
            for s in reversed(re.split(sentence_pattern, current_chunk)):
                if overlap_size + len(s) <= chunk_overlap:
                    overlap_text = s + " " + overlap_text
                    overlap_size += len(s) + 1
                else:
                    break
            
            current_chunk = overlap_text + sentence
        else:
            current_chunk += sentence + " "
    
    # Add the last chunk if it's not empty
    if current_chunk.strip():
        chunks.append(current_chunk)
    
    return chunks


def chunk_text_by_token(
    text: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 200
) -> List[str]:
    """
    Split text into chunks by approximate token count.
    
    Args:
        text: Text to split
        chunk_size: Maximum number of tokens per chunk
        chunk_overlap: Overlap between chunks in tokens
        
    Returns:
        List of text chunks
    """
    if not text:
        return []
    
    # Simple tokenization by splitting on whitespace
    # This is a rough approximation, not as accurate as a real tokenizer
    tokens = text.split()
    
    chunks = []
    start = 0
    
    while start < len(tokens):
        # Get chunk of size chunk_size
        end = min(start + chunk_size, len(tokens))
        
        # Create chunk from tokens
        chunk = " ".join(tokens[start:end])
        chunks.append(chunk)
        
        # Move start position for next chunk, with overlap
        start = end - chunk_overlap
        
        # Ensure start doesn't go negative
        start = max(0, start)
    
    return chunks


def chunk_text(
    text: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    chunk_method: Literal["token", "character", "sentence"] = "sentence"
) -> List[str]:
    """
    Split text into chunks using the specified method.
    
    Args:
        text: Text to split
        chunk_size: Maximum size of each chunk
        chunk_overlap: Overlap between chunks
        chunk_method: Method to use for chunking
        
    Returns:
        List of text chunks
    """
    if chunk_method == "token":
        return chunk_text_by_token(text, chunk_size, chunk_overlap)
    
    elif chunk_method == "sentence":
        return chunk_text_by_sentence(text, chunk_size, chunk_overlap)
    
    else:  # Default to character
        return chunk_text_by_character(text, chunk_size, chunk_overlap)
