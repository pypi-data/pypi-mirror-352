from AI import iNotes
import os
import pdfplumber
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

def summarize_notes(filepath="example_document", 
                    model="deepseek/deepseek-r1", 
                    format="docx", 
                    length="1000"):

    print(f"📂 Reading file: {filepath}.{format}")
    text_data = ""

    # Read the input file based on format
    try:
        if format == "txt":
            with open(filepath + ".txt", "r", encoding="utf-8") as f:
                text_data = f.read()

        elif format == "docx":
            doc = Document(filepath + ".docx")
            text_data = "\n".join([para.text for para in doc.paragraphs])

        elif format == "pdf":
            with pdfplumber.open(filepath + ".pdf") as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_data += page_text + "\n"

        else:
            raise ValueError("❌ Unsupported format. Choose from: pdf, txt, docx")
    except Exception as e:
        print(f"❌ Error while reading file: {e}")
        return

    if not text_data.strip():
        print("⚠️ File is empty or unreadable.")
        return

    print(f"🧠 Summarizing... (model: {model}, words: {length})")

    try:
        summarizer = iNotes(
            model=model,
            system_prompt=f"You are an AI notes summarizer (name: iNotes) which summarizes the given notes in {length} words. In the output summary, keep headings starting with ## and subheadings starting with #."
        )
        summary = summarizer.send_request(text_data)
    except Exception as e:
        print(f"❌ Error during summarization: {e}")
        return

    if not summary or not summary.strip():
        print("⚠️ No summary returned by the model.")
        return

    summary = summary.replace("\\n", "\n").replace("/n", "\n")
    base_name = os.path.splitext(os.path.basename(filepath))[0]
    output_path = f"{base_name}_summary"

    print("📝 Writing summary to file...")

    try:
        if format == "txt":
            output_file = output_path + ".txt"
            with open(output_file, "w", encoding="utf-8") as f:
                for paragraph in summary.split('\n'):
                    paragraph = paragraph.strip()
                    if not paragraph:
                        f.write("\n")
                        continue
                    if paragraph.startswith("##"):
                        f.write(f"\n**{paragraph.lstrip('#').strip().upper()}**\n")
                    elif paragraph.startswith("#"):
                        f.write(f"\n*{paragraph.lstrip('#').strip().title()}*\n")
                    else:
                        f.write(paragraph + "\n")

        elif format == "docx":
            output_file = output_path + ".docx"
            doc = Document()
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)

            for paragraph in summary.split('\n'):
                paragraph = paragraph.strip()
                if not paragraph:
                    doc.add_paragraph()
                    continue

                if paragraph.startswith("##"):
                    p = doc.add_paragraph()
                    run = p.add_run(paragraph.lstrip("#").strip())
                    run.bold = True
                    run.font.size = Pt(14)
                    run.font.name = 'Arial'
                elif paragraph.startswith("#"):
                    p = doc.add_paragraph()
                    run = p.add_run(paragraph.lstrip("#").strip())
                    run.bold = True
                    run.font.size = Pt(12)
                    run.font.name = 'Arial'
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(paragraph)
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'

            doc.save(output_file)

        elif format == "pdf":
            output_file = output_path + ".pdf"
            c = canvas.Canvas(output_file, pagesize=letter)
            width, height = letter
            margin = 50
            y = height - margin
            line_height = 16
            max_line_width = width - 2 * margin

            def wrap_text(text_line, font_name, font_size):
                words = text_line.split()
                lines = []
                current = ""
                for word in words:
                    test_line = current + " " + word if current else word
                    if c.stringWidth(test_line, font_name, font_size) <= max_line_width:
                        current = test_line
                    else:
                        lines.append(current)
                        current = word
                if current:
                    lines.append(current)
                return lines

            for paragraph in summary.split('\n'):
                paragraph = paragraph.strip()
                if not paragraph:
                    y -= line_height
                    continue

                if paragraph.startswith("##"):
                    font_name = "Helvetica-Bold"
                    font_size = 14
                    clean_text = paragraph.lstrip("#").strip()
                elif paragraph.startswith("#"):
                    font_name = "Helvetica-Bold"
                    font_size = 12
                    clean_text = paragraph.lstrip("#").strip()
                else:
                    font_name = "Helvetica"
                    font_size = 10
                    clean_text = paragraph

                c.setFont(font_name, font_size)
                lines = wrap_text(clean_text, font_name, font_size)
                for line in lines:
                    if y <= margin:
                        c.showPage()
                        y = height - margin
                        c.setFont(font_name, font_size)
                    c.drawString(margin, y, line)
                    y -= line_height

            c.save()

        else:
            raise ValueError("❌ Format not supported for saving.")

        print(f"✅ Summary saved as: {output_file}")

    except Exception as e:
        print(f"❌ Failed to save summary: {e}")



if __name__ == "__main__":
    summarize_notes("my_file", format="txt")
