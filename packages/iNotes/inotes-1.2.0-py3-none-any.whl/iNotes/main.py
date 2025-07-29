from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Pt
from AI import iNotes

def generate_notes(topic, 
                   filepath="output_notes", 
                   format="pdf",
                   model="deepseek/deepseek-r1",
                   short_notes = False):
    
    print("📌 Starting note generation...")
    print(f"📝 Topic: {topic}")
    print(f"📂 Output Path: {filepath}")
    print(f"📄 Format: {format}")
    print(f"🤖 Model: {model}")
    print(f"📏 Note length: {'Short' if short_notes else 'Long'}")

    if short_notes:
        filepath += "_short"
        system_prompt = "You are an AI notes maker (name: iNotes) which makes notes on the basis of given prompt. Make short notes and cover all the topics. Keep Headings startswith # and subheadings startswith ##."
    else:
        system_prompt = "You are an AI notes maker (name: iNotes) which makes notes on the basis of given prompt. Make long notes and cover all the topics. Keep Headings startswith # and subheadings startswith ##."

    print("📡 Sending request to AI model...")
    Notes = iNotes(system_prompt=system_prompt, model=model)
    text = Notes.send_request(topic)

    if text:
        print("📥 Response received. Processing text...")
        text = text.replace("\\n", "\n").replace("/n", "\n")

        if format == "pdf":
            print("🖨️ Generating PDF file...")
            filepath += ".pdf"
            c = canvas.Canvas(filepath, pagesize=letter)
            width, height = letter
            margin = 50
            y_position = height - margin
            line_height = 16
            max_line_width = width - 2 * margin

            def split_text_by_width(text_line, font_name, font_size):
                words = text_line.split()
                lines = []
                current_line = ""
                for word in words:
                    test_line = current_line + " " + word if current_line else word
                    if c.stringWidth(test_line, font_name, font_size) <= max_line_width:
                        current_line = test_line
                    else:
                        lines.append(current_line)
                        current_line = word
                if current_line:
                    lines.append(current_line)
                return lines

            for paragraph in text.split('\n'):
                paragraph = paragraph.strip()
                if not paragraph:
                    y_position -= line_height
                    continue

                if paragraph.startswith("#"):
                    font_name = "Helvetica-Bold"
                    font_size = 14
                    clean_text = paragraph.lstrip("#").strip()
                elif paragraph.startswith("##"):
                    font_name = "Helvetica-Bold"
                    font_size = 12
                    clean_text = paragraph.lstrip("#").strip()
                else:
                    font_name = "Helvetica"
                    font_size = 10
                    clean_text = paragraph

                c.setFont(font_name, font_size)
                lines = split_text_by_width(clean_text, font_name, font_size)
                for line in lines:
                    if y_position <= margin:
                        c.showPage()
                        y_position = height - margin
                        c.setFont(font_name, font_size)
                    c.drawString(margin, y_position, line)
                    y_position -= line_height

            c.save()

        elif format == "txt":
            print("🖊️ Generating TXT file...")
            filepath += ".txt"
            with open(filepath, "w", encoding="utf-8") as f:
                for paragraph in text.split('\n'):
                    paragraph = paragraph.strip()
                    if not paragraph:
                        f.write("\n")
                        continue
                    if paragraph.startswith("#"):
                        f.write(f"\n**{paragraph.lstrip('#').strip().upper()}**\n")
                    elif paragraph.startswith("##"):
                        f.write(f"\n*{paragraph.lstrip('#').strip().title()}*\n")
                    else:
                        f.write(paragraph + "\n")

        elif format == "docx":
            print("📝 Generating DOCX file...")
            filepath += ".docx"
            doc = Document()
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)

            for paragraph in text.split('\n'):
                paragraph = paragraph.strip()
                if not paragraph:
                    doc.add_paragraph()
                    continue

                if paragraph.startswith("#"):
                    p = doc.add_paragraph()
                    run = p.add_run(paragraph.lstrip("#").strip())
                    run.bold = True
                    run.font.size = Pt(14)
                    run.font.name = 'Arial'
                elif paragraph.startswith("##"):
                    p = doc.add_paragraph()
                    run = p.add_run(paragraph.lstrip("#").strip())
                    run.bold = True
                    run.font.size = Pt(12)
                    run.font.name = 'Arial'
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(paragraph)
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'

            doc.save(filepath)

        else:
            print("❌ Unsupported file format selected.")
            raise ValueError("❌ Unsupported format. Choose from: pdf, txt, docx")

        print(f"✅ Notes successfully saved to: {filepath}")
    else:
        print("⚠️ No response received from the model.")

# Example Usage
if __name__ == "__main__":
    while True:
        prompt = input("Enter a topic for notes :\n>>>> ")

        if "Error" not in prompt:
            print("🚀 Generating notes...")
            generate_notes(prompt, filepath=f"{prompt.replace(' ', '_')}", format="docx", short_notes=False)
            print("✅ Success!\n")
