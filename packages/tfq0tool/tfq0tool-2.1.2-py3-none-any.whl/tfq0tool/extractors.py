"""Text extraction implementations for various file formats."""

import os
import logging
from typing import Optional, Dict, Any
from pathlib import Path
import chardet
from PyPDF2 import PdfReader
from docx import Document
import openpyxl
from PIL import Image
import pytesseract
from pdfminer.high_level import extract_text as pdfminer_extract
from pdfminer.pdfparser import PDFSyntaxError

logger = logging.getLogger(__name__)

class BaseExtractor:
    """Base class for all text extractors."""
    
    def __init__(self, preprocessing_options: Optional[Dict[str, Any]] = None):
        self.preprocessing_options = preprocessing_options or {}
    
    def detect_encoding(self, file_path: str) -> str:
        """Detect the encoding of a file."""
        with open(file_path, 'rb') as f:
            raw_data = f.read()
            result = chardet.detect(raw_data)
            return result['encoding'] or 'utf-8'
    
    def preprocess_text(self, text: str) -> str:
        """Apply text preprocessing based on options."""
        if self.preprocessing_options.get('strip_whitespace', True):
            text = ' '.join(text.split())
        if self.preprocessing_options.get('lowercase', False):
            text = text.lower()
        return text

class TxtExtractor(BaseExtractor):
    def extract(self, file_path: str) -> str:
        encoding = self.detect_encoding(file_path)
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                text = f.read()
            return self.preprocess_text(text)
        except UnicodeDecodeError as e:
            logger.error(f"Error decoding {file_path}: {e}")
            raise

class PdfExtractor(BaseExtractor):
    def extract(self, file_path: str, password: Optional[str] = None) -> str:
        text_parts = []
        try:
            # Try PyPDF2 first
            with open(file_path, 'rb') as f:
                reader = PdfReader(f)
                if reader.is_encrypted and password:
                    reader.decrypt(password)
                
                for page in reader.pages:
                    text = page.extract_text()
                    if not text.strip():
                        # If no text found, try pdfminer
                        text = pdfminer_extract(file_path, page_numbers=[reader.pages.index(page)])
                    if not text.strip():
                        # If still no text, might be scanned PDF - try OCR
                        text = self._extract_text_with_ocr(file_path, reader.pages.index(page))
                    text_parts.append(text)
                
            return self.preprocess_text('\n'.join(text_parts))
        except (PDFSyntaxError, Exception) as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise

    def _extract_text_with_ocr(self, file_path: str, page_number: int) -> str:
        """Extract text from images using OCR."""
        try:
            # Convert PDF page to image and perform OCR
            images = self._pdf_page_to_image(file_path, page_number)
            text_parts = []
            for img in images:
                text_parts.append(pytesseract.image_to_string(img))
            return '\n'.join(text_parts)
        except Exception as e:
            logger.warning(f"OCR extraction failed for page {page_number}: {e}")
            return ""

    def _pdf_page_to_image(self, file_path: str, page_number: int):
        """Convert a PDF page to image."""
        # Implementation depends on your preferred PDF to image conversion method
        # This is a placeholder - you would need to implement actual conversion
        pass

class DocxExtractor(BaseExtractor):
    def extract(self, file_path: str) -> str:
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract headers
            for section in doc.sections:
                header = section.header
                if header.is_linked_to_previous:
                    continue
                for paragraph in header.paragraphs:
                    text_parts.append(paragraph.text)
            
            # Extract main content
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' '.join(cell.text for cell in row.cells)
                    text_parts.append(row_text)
            
            # Extract footers
            for section in doc.sections:
                footer = section.footer
                if footer.is_linked_to_previous:
                    continue
                for paragraph in footer.paragraphs:
                    text_parts.append(paragraph.text)
            
            return self.preprocess_text('\n'.join(text_parts))
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise

class XlsxExtractor(BaseExtractor):
    def extract(self, file_path: str) -> str:
        try:
            text_parts = []
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            for sheet in workbook.sheetnames:
                worksheet = workbook[sheet]
                text_parts.append(f"\n=== Sheet: {sheet} ===\n")
                
                # Get the maximum row and column with content
                max_row = worksheet.max_row
                max_col = worksheet.max_column
                
                for row in worksheet.iter_rows(values_only=True, max_row=max_row, max_col=max_col):
                    row_text = " | ".join(str(cell) if cell is not None else '' for cell in row)
                    if row_text.strip():  # Only add non-empty rows
                        text_parts.append(row_text)
            
            return self.preprocess_text('\n'.join(text_parts))
        except Exception as e:
            logger.error(f"Error processing XLSX {file_path}: {e}")
            raise

class CodeExtractor(BaseExtractor):
    def extract(self, file_path: str) -> str:
        encoding = self.detect_encoding(file_path)
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                text = f.read()
            return self.preprocess_text(text)
        except UnicodeDecodeError as e:
            logger.error(f"Error decoding {file_path}: {e}")
            raise

def get_extractor(file_path: str, preprocessing_options: Optional[Dict[str, Any]] = None) -> BaseExtractor:
    """Factory function to get the appropriate extractor based on file extension."""
    ext = Path(file_path).suffix.lower()
    
    extractors = {
        '.txt': TxtExtractor,
        '.pdf': PdfExtractor,
        '.docx': DocxExtractor,
        '.xlsx': XlsxExtractor,
        '.py': CodeExtractor,
        '.java': CodeExtractor,
        '.js': CodeExtractor,
        '.html': CodeExtractor,
        '.css': CodeExtractor,
        '.json': CodeExtractor,
        '.xml': CodeExtractor,
        '.c': CodeExtractor,
        '.cpp': CodeExtractor,
    }
    
    extractor_class = extractors.get(ext)
    if not extractor_class:
        raise ValueError(f"Unsupported file format: {ext}")
    
    return extractor_class(preprocessing_options) 