"""Text extraction module with support for various file formats."""

import logging
from pathlib import Path
from typing import Dict, Any, Optional, Generator, Union
from abc import ABC, abstractmethod
import mimetypes
import tempfile
import shutil

# Third-party imports for different file formats
from PyPDF2 import PdfReader
from docx import Document
import pandas as pd
import xml.etree.ElementTree as ET
import json
import chardet

from .config import config

logger = logging.getLogger(__name__)

class BaseExtractor(ABC):
    """Base class for text extractors."""
    
    def __init__(self, preprocessing_options: Optional[Dict[str, Any]] = None):
        self.preprocessing_options = preprocessing_options or {}
    
    @abstractmethod
    def extract_iter(self, file_path: str) -> Generator[str, None, None]:
        """Extract text from file in chunks."""
        pass
    
    def extract(self, file_path: str) -> str:
        """Extract all text at once. Use with caution for large files."""
        return "".join(self.extract_iter(file_path))
    
    def preprocess_text(self, text: str) -> str:
        """Apply preprocessing options to text."""
        if self.preprocessing_options.get("lowercase", False):
            text = text.lower()
        if self.preprocessing_options.get("strip_whitespace", False):
            text = " ".join(text.split())
        return text

class PDFExtractor(BaseExtractor):
    """Extract text from PDF files."""
    
    def extract_iter(self, file_path: str) -> Generator[str, None, None]:
        try:
            with open(file_path, 'rb') as file:
                pdf = PdfReader(file)
                
                if pdf.is_encrypted:
                    password = self.preprocessing_options.get("password")
                    if not password or not pdf.decrypt(password):
                        raise ValueError("PDF is encrypted and no valid password provided")
                
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        yield self.preprocess_text(text)
                        
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            raise

class DocxExtractor(BaseExtractor):
    """Extract text from DOCX files."""
    
    def extract_iter(self, file_path: str) -> Generator[str, None, None]:
        try:
            doc = Document(file_path)
            
            # Process paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text
                if text:
                    yield self.preprocess_text(text + "\n")
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text:
                        yield self.preprocess_text(row_text + "\n")
                        
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise

class SpreadsheetExtractor(BaseExtractor):
    """Extract text from Excel and CSV files."""
    
    def extract_iter(self, file_path: str) -> Generator[str, None, None]:
        try:
            # Determine file type
            if file_path.endswith('.csv'):
                df = pd.read_csv(file_path, chunksize=1000)
            else:  # Excel files
                df = pd.read_excel(file_path, sheet_name=None)
                if isinstance(df, dict):
                    for sheet_name, sheet_df in df.items():
                        yield f"\n=== Sheet: {sheet_name} ===\n"
                        for chunk in self._process_dataframe(sheet_df):
                            yield chunk
                    return
            
            # Process chunks
            for chunk in self._process_dataframe(df):
                yield chunk
                
        except Exception as e:
            logger.error(f"Error extracting text from spreadsheet {file_path}: {e}")
            raise
    
    def _process_dataframe(self, df) -> Generator[str, None, None]:
        """Process DataFrame in chunks."""
        if hasattr(df, 'chunksize'):  # CSV with chunks
            for chunk in df:
                yield self.preprocess_text(chunk.to_string() + "\n")
        else:  # Single DataFrame
            yield self.preprocess_text(df.to_string() + "\n")

class XMLExtractor(BaseExtractor):
    """Extract text from XML files."""
    
    def extract_iter(self, file_path: str) -> Generator[str, None, None]:
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            for elem in root.iter():
                if elem.text and elem.text.strip():
                    yield self.preprocess_text(elem.text.strip() + "\n")
                if elem.tail and elem.tail.strip():
                    yield self.preprocess_text(elem.tail.strip() + "\n")
                    
        except Exception as e:
            logger.error(f"Error extracting text from XML {file_path}: {e}")
            raise

class JSONExtractor(BaseExtractor):
    """Extract text from JSON files."""
    
    def extract_iter(self, file_path: str) -> Generator[str, None, None]:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            for text in self._extract_values(data):
                if text:
                    yield self.preprocess_text(str(text) + "\n")
                    
        except Exception as e:
            logger.error(f"Error extracting text from JSON {file_path}: {e}")
            raise
    
    def _extract_values(self, obj: Any) -> Generator[str, None, None]:
        """Recursively extract values from JSON object."""
        if isinstance(obj, dict):
            for value in obj.values():
                yield from self._extract_values(value)
        elif isinstance(obj, list):
            for item in obj:
                yield from self._extract_values(item)
        elif isinstance(obj, (str, int, float, bool)):
            yield str(obj)

class TextExtractor(BaseExtractor):
    """Extract text from plain text files with encoding detection."""
    
    def extract_iter(self, file_path: str) -> Generator[str, None, None]:
        try:
            # Detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                result = chardet.detect(raw_data)
                encoding = result['encoding'] or 'utf-8'
            
            # Read file in chunks
            chunk_size = config.get("processing", "chunk_size")
            with open(file_path, 'r', encoding=encoding) as f:
                while True:
                    chunk = f.read(chunk_size)
                    if not chunk:
                        break
                    yield self.preprocess_text(chunk)
                    
        except Exception as e:
            logger.error(f"Error extracting text from file {file_path}: {e}")
            raise

def get_extractor(file_path: str, preprocessing_options: Optional[Dict[str, Any]] = None) -> BaseExtractor:
    """Factory function to get appropriate extractor based on file type."""
    file_path = str(file_path).lower()
    
    extractors = {
        '.pdf': PDFExtractor,
        '.docx': DocxExtractor,
        '.doc': DocxExtractor,
        '.xlsx': SpreadsheetExtractor,
        '.xls': SpreadsheetExtractor,
        '.csv': SpreadsheetExtractor,
        '.xml': XMLExtractor,
        '.json': JSONExtractor,
        '.txt': TextExtractor,
        '.log': TextExtractor,
        '.md': TextExtractor,
        '.rtf': TextExtractor
    }
    
    ext = Path(file_path).suffix.lower()
    extractor_class = extractors.get(ext)
    
    if not extractor_class:
        mime_type = mimetypes.guess_type(file_path)[0]
        if mime_type and 'text' in mime_type:
            extractor_class = TextExtractor
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    return extractor_class(preprocessing_options) 