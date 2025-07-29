import os
import shutil
import sys
import tempfile
import zipfile

import docx
from docx import Document
from docx.image.exceptions import UnrecognizedImageError
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt as DocxPt

from chgksuite.common import log_wrap, replace_escaped
from chgksuite.composer.composer_common import BaseExporter, backtick_replace, parseimg

WHITEN = {
    "handout": False,
    "zachet": True,
    "nezachet": True,
    "comment": True,
    "source": True,
    "author": False,
}


def replace_font_in_docx(template_path, new_font):
    """Replace Arial fonts with specified font in docx template"""
    temp_dir = tempfile.mkdtemp()
    template_name = os.path.basename(template_path)
    temp_template = os.path.join(temp_dir, template_name)
    shutil.copy2(template_path, temp_template)

    temp_zip = os.path.join(temp_dir, "template.zip")
    os.rename(temp_template, temp_zip)
    with zipfile.ZipFile(temp_zip, "r") as zip_ref:
        zip_ref.extractall(temp_dir)
    os.remove(temp_zip)

    for root, _, files in os.walk(temp_dir):
        for file in files:
            if file.endswith(".xml"):
                file_path = os.path.join(root, file)
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        content = f.read()

                    content = content.replace("Arial Unicode MS", new_font)
                    content = content.replace("Arial", new_font)

                    with open(file_path, "w", encoding="utf-8") as f:
                        f.write(content)
                except UnicodeError:
                    continue

    shutil.make_archive(temp_template, "zip", temp_dir)
    os.rename(temp_template + ".zip", temp_template)
    return temp_template


class DocxExporter(BaseExporter):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.qcount = 0

        if self.args.font_face:
            self.args.docx_template = replace_font_in_docx(
                self.args.docx_template, self.args.font_face
            )

    def __del__(self):
        # Cleanup temp directory if it exists
        if hasattr(self, "temp_dir") and os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def _docx_format(self, *args, **kwargs):
        kwargs.update(self.dir_kwargs)
        return self.docx_format(*args, **kwargs)

    def docx_format(self, el, para, whiten, **kwargs):
        if isinstance(el, list):
            if len(el) > 1 and isinstance(el[1], list):
                self.docx_format(el[0], para, whiten, **kwargs)
                licount = 0
                for li in el[1]:
                    licount += 1

                    para.add_run("\n{}. ".format(licount))
                    self.docx_format(li, para, whiten, **kwargs)
            else:
                licount = 0
                for li in el:
                    licount += 1

                    para.add_run("\n{}. ".format(licount))
                    self.docx_format(li, para, whiten, **kwargs)

        if isinstance(el, str):
            self.logger.debug("parsing element {}:".format(log_wrap(el)))

            if kwargs.get("remove_accents"):
                el = el.replace("\u0301", "")
            if kwargs.get("remove_brackets"):
                el = self.remove_square_brackets(el)
            else:
                el = replace_escaped(el)

            el = backtick_replace(el)

            for run in self.parse_4s_elem(el):
                if run[0] == "pagebreak":
                    if self.args.spoilers == "dots":
                        for _ in range(30):
                            para = self.doc.add_paragraph()
                            para.add_run(".")
                        para = self.doc.add_paragraph()
                    else:
                        para = self.doc.add_page_break()
                elif run[0] == "linebreak":
                    para.add_run("\n")
                elif run[0] == "screen":
                    if kwargs.get("remove_accents") or kwargs.get("remove_brackets"):
                        text = run[1]["for_screen"]
                    else:
                        text = run[1]["for_print"]
                    if kwargs.get("replace_no_break_spaces"):
                        text = self._replace_no_break(text)
                    r = para.add_run(text)
                elif run[0] == "hyperlink" and not (
                    whiten and self.args.spoilers == "whiten"
                ):
                    r = self.add_hyperlink(para, run[1], run[1])
                elif run[0] == "img":
                    if run[1].endswith(".shtml"):
                        r = para.add_run(
                            "(ТУТ БЫЛА ССЫЛКА НА ПРОТУХШУЮ КАРТИНКУ)\n"
                        )  # TODO: добавить возможность пропускать кривые картинки опцией
                        continue
                    parsed_image = parseimg(
                        run[1],
                        dimensions="inches",
                        tmp_dir=kwargs.get("tmp_dir"),
                        targetdir=kwargs.get("targetdir"),
                    )
                    imgfile = parsed_image["imgfile"]
                    width = parsed_image["width"]
                    height = parsed_image["height"]
                    inline = parsed_image["inline"]
                    if inline:
                        r = para.add_run("")
                    else:
                        r = para.add_run("\n")

                    try:
                        if inline:
                            r.add_picture(
                                imgfile,
                                height=Inches(
                                    1.0 / 6
                                ),  # Height is based on docx template
                            )
                        else:
                            r.add_picture(
                                imgfile, width=Inches(width), height=Inches(height)
                            )
                    except UnrecognizedImageError:
                        sys.stderr.write(
                            f"python-docx can't recognize header for {imgfile}\n"
                        )
                    if not inline:
                        r = para.add_run("\n")
                    continue
                else:
                    text = run[1]
                    if kwargs.get("replace_no_break_spaces"):
                        text = self._replace_no_break(text)
                    r = para.add_run(text)
                    if "italic" in run[0]:
                        r.italic = True
                    if "bold" in run[0]:
                        r.bold = True
                    if "underline" in run[0]:
                        r.underline = True
                    if run[0] == "strike":
                        r.font.strike = True
                    if run[0] == "sc":
                        r.small_caps = True
                    if whiten and self.args.spoilers == "whiten":
                        r.style = "Whitened"

    def add_hyperlink(self, paragraph, text, url):
        # adapted from https://github.com/python-openxml/python-docx/issues/610
        doc = self.doc
        run = paragraph.add_run(text)
        run.style = doc.styles["Hyperlink"]
        part = paragraph.part
        r_id = part.relate_to(
            url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
        )
        hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
        hyperlink.set(docx.oxml.shared.qn("r:id"), r_id)
        hyperlink.append(run._r)
        paragraph._p.append(hyperlink)
        return hyperlink

    def add_question(
        self, element, skip_qcount=False, screen_mode=False, external_para=None
    ):
        q = element[1]
        if external_para is None:
            p = self.doc.add_paragraph()
        else:
            p = external_para
        p.paragraph_format.space_before = DocxPt(18)
        p.paragraph_format.keep_together = True
        if "number" not in q and not skip_qcount:
            self.qcount += 1
        if "setcounter" in q:
            self.qcount = int(q["setcounter"])
        p.add_run(
            "{question}. ".format(
                question=self.get_label(
                    q,
                    "question",
                    number=self.qcount if "number" not in q else q["number"],
                )
            )
        ).bold = True

        if "handout" in q:
            p.add_run("\n[{handout}: ".format(handout=self.get_label(q, "handout")))
            self._docx_format(
                q["handout"],
                p,
                WHITEN["handout"],
                remove_accents=screen_mode,
                remove_brackets=screen_mode,
            )
            p.add_run("\n]")
        if not self.args.noparagraph:
            p.add_run("\n")

        self._docx_format(
            q["question"],
            p,
            False,
            remove_accents=screen_mode,
            remove_brackets=screen_mode,
            replace_no_break_spaces=True,
        )

        if not self.args.noanswers:
            if self.args.spoilers == "pagebreak":
                p = self.doc.add_page_break()
            elif self.args.spoilers == "dots":
                for _ in range(30):
                    if external_para is None:
                        p = self.doc.add_paragraph()
                    else:
                        p.add_run("\n")
                    p.add_run(".")
                if external_para is None:
                    p = self.doc.add_paragraph()
                else:
                    p.add_run("\n")
            else:
                if external_para is None:
                    p = self.doc.add_paragraph()
                else:
                    p.add_run("\n")
            p.paragraph_format.keep_together = True
            p.paragraph_format.space_before = DocxPt(6)
            p.add_run(f"{self.get_label(q, 'answer')}: ").bold = True
            self._docx_format(
                q["answer"],
                p,
                True,
                remove_accents=screen_mode,
                replace_no_break_spaces=True,
            )

            for field in ["zachet", "nezachet", "comment", "source", "author"]:
                if field in q:
                    if field == "source":
                        if external_para is None:
                            p = self.doc.add_paragraph()
                            p.paragraph_format.keep_together = True
                        else:
                            p.add_run("\n")
                    else:
                        p.add_run("\n")
                    p.add_run(f"{self.get_label(q, field)}: ").bold = True
                    self._docx_format(
                        q[field],
                        p,
                        WHITEN[field],
                        remove_accents=screen_mode,
                        remove_brackets=screen_mode,
                        replace_no_break_spaces=field != "source",
                    )

    def _add_question_columns(self, element):
        table = self.doc.add_table(rows=1, cols=2)
        table.autofit = True

        def set_cell_border(cell):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            for edge in ["top", "left", "bottom", "right"]:
                border = OxmlElement("w:{}Border".format(edge))
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "auto")
                tcPr.append(border)

        for row in table.rows:
            for cell in row.cells:
                set_cell_border(cell)

        table.cell(0, 0).paragraphs[0].add_run("Версия для ведущего\n").bold = True
        table.cell(0, 1).paragraphs[0].add_run("Версия для экрана\n").bold = True

        self.add_question(
            element, screen_mode=False, external_para=table.cell(0, 0).paragraphs[0]
        )
        self.add_question(
            element, screen_mode=True, external_para=table.cell(0, 1).paragraphs[0]
        )

        self.doc.add_paragraph()

    def _add_question_content(self, q, p, skip_qcount=False, screen_mode=False):
        """Helper method to add question content to a paragraph"""
        if "number" not in q and not skip_qcount:
            self.qcount += 1
        if "setcounter" in q:
            self.qcount = int(q["setcounter"])
        p.add_run(
            "{question}. ".format(
                question=self.get_label(
                    q,
                    "question",
                    number=self.qcount if "number" not in q else q["number"],
                )
            )
        ).bold = True

        if "handout" in q:
            p.add_run("\n[{handout}: ".format(handout=self.get_label(q, "handout")))
            self._docx_format(
                q["handout"],
                p,
                WHITEN["handout"],
                remove_accents=screen_mode,
                remove_brackets=screen_mode,
            )
            p.add_run("\n]")
        if not self.args.noparagraph:
            p.add_run("\n")

        self._docx_format(
            q["question"],
            p,
            False,
            remove_accents=screen_mode,
            remove_brackets=screen_mode,
            replace_no_break_spaces=True,
        )

        if not self.args.noanswers:
            if self.args.spoilers == "pagebreak":
                p = self.doc.add_page_break()
            elif self.args.spoilers == "dots":
                for _ in range(30):
                    p = self.doc.add_paragraph()
                    p.add_run(".")
                p = self.doc.add_paragraph()
            else:
                p = self.doc.add_paragraph()
            p.paragraph_format.keep_together = True
            p.paragraph_format.space_before = DocxPt(6)
            p.add_run(f"{self.get_label(q, 'answer')}: ").bold = True
            self._docx_format(
                q["answer"],
                p,
                True,
                remove_accents=screen_mode,
                replace_no_break_spaces=True,
            )

            for field in ["zachet", "nezachet", "comment", "source", "author"]:
                if field in q:
                    if field == "source":
                        p = self.doc.add_paragraph()
                        p.paragraph_format.keep_together = True
                    else:
                        p.add_run("\n")
                    p.add_run(f"{self.get_label(q, field)}: ").bold = True
                    self._docx_format(
                        q[field],
                        p,
                        WHITEN[field],
                        remove_accents=screen_mode,
                        remove_brackets=screen_mode,
                        replace_no_break_spaces=field != "source",
                    )

    def export(self, outfilename):
        self.logger.debug(self.args.docx_template)
        self.doc = Document(self.args.docx_template)
        self.logger.debug(log_wrap(self.structure))

        firsttour = True
        prev_element = None
        para = None
        page_break_before_heading = False
        for element in self.structure:
            if element[0] == "meta":
                para = self.doc.add_paragraph()
                if prev_element and prev_element[0] == "Question":
                    para.paragraph_format.space_before = DocxPt(18)
                self._docx_format(element[1], para, False, replace_no_break_spaces=True)
                self.doc.add_paragraph()

            if element[0] in ["editor", "date", "heading", "section"]:
                if element[0] == "heading" and para is not None:
                    page_break_before_heading = True
                if para is None:
                    para = self.doc.paragraphs[0]
                else:
                    para = self.doc.add_paragraph()
                self._docx_format(element[1], para, False, replace_no_break_spaces=True)
                if element[0] == "heading" and page_break_before_heading:
                    para.paragraph_format.page_break_before = True
                if element[0] == "section":
                    if not firsttour:
                        para.paragraph_format.page_break_before = True
                    else:
                        firsttour = False
                if element[0] == "heading":
                    for st in self.doc.styles:
                        if st.name == "Heading 1":
                            break
                    para.style = st
                elif element[0] == "section":
                    for st in self.doc.styles:
                        if st.name == "Heading 2":
                            break
                    para.style = st
                para.paragraph_format.keep_with_next = True
                para.add_run("\n")

            if element[0] == "Question":
                if self.args.screen_mode == "add_versions_columns":
                    self._add_question_columns(element)
                elif self.args.screen_mode == "add_versions":
                    para = self.doc.add_paragraph()
                    para = self.doc.add_paragraph()
                    para.add_run("Версия для ведущего:").bold = True
                    self.add_question(element, screen_mode=False)
                    para = self.doc.add_paragraph()
                    para = self.doc.add_paragraph()
                    para.add_run("Версия для экрана:").bold = True
                    self.add_question(element, skip_qcount=True, screen_mode=True)
                elif self.args.screen_mode == "replace_all":
                    self.add_question(element, screen_mode=True)
                else:
                    self.add_question(element)
            prev_element = element

        self.doc.save(outfilename)
        self.logger.info("Output: {}".format(outfilename))
